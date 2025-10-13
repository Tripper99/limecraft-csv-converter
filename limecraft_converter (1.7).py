#!/usr/bin/env python3
"""
Limecraft CSV-konverterare
Konverterar Limecraft transkription CSV-filer till Word-dokument och Inqscribe-filer samt justerar tidskoder.
Av Dan Josefsson 2025-01-08
Version 1.7
"""

import csv
import math
from pathlib import Path
import sys
from typing import Dict, List, Optional

try:
    import ttkbootstrap as tb
    from ttkbootstrap.constants import DANGER, INFO, PRIMARY, SUCCESS
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext
except ImportError:
    print("Error: ttkbootstrap not installed. Install with: pip install ttkbootstrap")
    sys.exit(1)

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Constants
FRAME_RATE = 30
TIME_THRESHOLD = 0.01
VERSION = "1.7"
REQUIRED_COLUMNS = ['Media Start', 'Transcript', 'Speakers']  # Added Speakers column
SUPPORTED_ENCODINGS = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']


# Custom Exception Classes
class LimecraftConverterError(Exception):
    """Base exception for all Limecraft Converter errors"""
    pass


class CSVValidationError(LimecraftConverterError):
    """Raised when CSV file validation fails"""
    pass


class TimecodeFormatError(LimecraftConverterError):
    """Raised when timecode format is invalid"""
    pass


class FileProcessingError(LimecraftConverterError):
    """Raised when file reading/writing fails"""
    pass


class ConversionError(LimecraftConverterError):
    """Raised when document conversion fails"""
    pass


class LimecraftConverter:
    def __init__(self):
        self.csv_file: Optional[str] = None
        self.data: List[Dict[str, str]] = []

    def normalize_time_input(self, time_str: str) -> str:
        """Normalize different time input formats to HH:MM:SS:FF"""
        if not time_str.strip():
            return ""

        time_str = time_str.strip()

        # Remove all non-digit characters and collect digits
        digits_only = ''.join(filter(str.isdigit, time_str))

        # Ensure we have exactly 8 digits (pad with zeros if needed)
        if len(digits_only) < 8:
            digits_only = digits_only.zfill(8)
        elif len(digits_only) > 8:
            digits_only = digits_only[:8]  # Take first 8 digits

        # Split into components
        hours = digits_only[:2]
        minutes = digits_only[2:4]
        seconds = digits_only[4:6]
        frames = digits_only[6:8]

        # Validate ranges
        if int(minutes) >= 60 or int(seconds) >= 60 or int(frames) >= FRAME_RATE:
            raise TimecodeFormatError(f"Invalid time values: minutes/seconds must be < 60, frames must be < {FRAME_RATE}")

        return f"{hours}:{minutes}:{seconds}:{frames}"

    def parse_timecode(self, timecode_str: str) -> float:
        """Parse timecode string to seconds"""
        if not timecode_str:
            return 0
        try:
            # Clean the timecode string
            timecode_str = str(timecode_str).strip()
            # Handle format: 00:00:03:08 or 00:00:03.08
            timecode_str = timecode_str.replace('.', ':')
            parts = timecode_str.split(':')
            if len(parts) == 4:
                hours, minutes, seconds, frames = map(int, parts)
                # Convert to total seconds
                total_seconds = hours * 3600 + minutes * 60 + seconds + frames / FRAME_RATE
                return total_seconds
            elif len(parts) == 3:
                hours, minutes, seconds = map(int, parts)
                return hours * 3600 + minutes * 60 + seconds
            elif len(parts) == 2:
                minutes, seconds = map(int, parts)
                return minutes * 60 + seconds
        except (ValueError, IndexError, TypeError):
            return 0
        return 0

    def seconds_to_timecode(self, seconds: float) -> str:
        """Convert seconds to timecode format [HH:MM:SS.FF]"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        frames = int((seconds % 1) * FRAME_RATE)
        return f"[{hours:02d}:{minutes:02d}:{secs:02d}.{frames:02d}]"

    def adjust_timecodes(self, start_time_str: str) -> None:
        """Add the specified time to all timecodes in the transcription"""
        if not start_time_str.strip():
            return  # No adjustment needed

        try:
            # Normalize the input time format
            normalized_time = self.normalize_time_input(start_time_str)

            # Parse the time to add
            time_to_add_seconds = self.parse_timecode(normalized_time)

            # Add this time to all existing timecodes
            for row in self.data:
                current_timecode = row.get('Media Start', '')
                current_seconds = self.parse_timecode(current_timecode)
                new_seconds = current_seconds + time_to_add_seconds
                row['Media Start'] = self.seconds_to_timecode(new_seconds).strip('[]')

        except Exception as e:
            raise TimecodeFormatError(f"Invalid time format: {str(e)}")

    def add_starting_timecode(self, time_adjustment_applied: bool = False) -> None:
        """Add [00:00:00.00] at the very beginning if transcript doesn't start at 00:00:00.00"""
        if not self.data:
            return

        # Check if first timecode is already 00:00:00.00
        first_timecode = self.data[0]['Media Start']
        first_seconds = self.parse_timecode(first_timecode)

        # Use math.isclose for robust floating point comparison
        # Always add starting timecode if first one is not at zero
        if not math.isclose(first_seconds, 0.0, abs_tol=TIME_THRESHOLD):
            # Add explanatory text if time was adjusted by user
            transcript_text = "(Starttid justerad av användare)" if time_adjustment_applied else ""

            starting_entry = {
                'Media Start': '00:00:00.00',
                'Transcript': transcript_text,
                'Speakers': ''  # Empty speaker for starting entry
            }
            # Insert at the beginning
            self.data.insert(0, starting_entry)

    def _clean_column_names(self, raw_data: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Clean column names by removing BOM and extra spaces"""
        cleaned_data = []
        for row in raw_data:
            cleaned_row = {}
            for key, value in row.items():
                # Clean the key: remove BOM, strip whitespace
                clean_key = key.strip().lstrip('\ufeff').strip()
                cleaned_row[clean_key] = value
            cleaned_data.append(cleaned_row)
        return cleaned_data

    def _parse_limecraft_combined_format(self, data: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Parse Limecraft CSV format where all data is in a single column"""
        normalized_data = []

        for row in data:
            # Get the first non-empty value from the row (should be the combined data)
            combined_data = None
            for value in row.values():
                if value and value.strip():
                    combined_data = value.strip()
                    break

            if not combined_data:
                continue

            # Split the combined data: "timecode,speaker,transcript"
            # Handle quoted content properly
            parts = []
            current_part = ""
            in_quotes = False
            quote_char = None

            i = 0
            while i < len(combined_data):
                char = combined_data[i]

                if char in ['"', "'"] and not in_quotes:
                    in_quotes = True
                    quote_char = char
                elif char == quote_char and in_quotes:
                    # Check for escaped quotes (double quotes)
                    if i + 1 < len(combined_data) and combined_data[i + 1] == quote_char:
                        current_part += char
                        i += 1  # Skip the next quote
                    else:
                        in_quotes = False
                        quote_char = None
                elif char == ',' and not in_quotes:
                    parts.append(current_part.strip())
                    current_part = ""
                    i += 1
                    continue
                else:
                    current_part += char

                i += 1

            # Add the last part
            if current_part.strip():
                parts.append(current_part.strip())

            # Ensure we have at least 3 parts
            while len(parts) < 3:
                parts.append("")

            # Clean up the parts (remove extra quotes)
            timecode = parts[0].strip('"\'')
            speaker = parts[1].strip('"\'')
            transcript = parts[2].strip('"\'')

            # Handle double quotes in transcript (Limecraft uses "" for quotes)
            transcript = transcript.replace('""', '"')

            normalized_row = {
                'Media Start': timecode,
                'Speakers': speaker,
                'Transcript': transcript
            }
            normalized_data.append(normalized_row)

        return normalized_data

    def _validate_and_normalize_columns(self, data: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """Validate required columns exist and keep only the columns we need"""
        if not data:
            raise CSVValidationError("CSV file is empty")

        actual_columns = list(data[0].keys())

        # Check if this is Limecraft's combined format (all data in one column)
        # This happens when the CSV has headers but data is quoted as a single field
        first_row_data = data[0]
        combined_format = False

        # Look for a single column containing comma-separated timecode data
        for value in first_row_data.values():
            if value and isinstance(value, str) and ':' in value and ',' in value:
                # Looks like timecode,speaker,transcript format
                combined_format = True
                break

        if combined_format:
            return self._parse_limecraft_combined_format(data)

        # Standard format - check if required columns exist (case-insensitive)
        missing_columns = []
        column_mapping = {}

        for req_col in REQUIRED_COLUMNS:
            found = False
            for actual_col in actual_columns:
                if req_col.lower() == actual_col.lower().strip():
                    column_mapping[req_col] = actual_col
                    found = True
                    break
            if not found:
                missing_columns.append(req_col)

        if missing_columns:
            raise CSVValidationError(f"Missing required columns: {missing_columns}. Found columns: {actual_columns}")

        # Keep ONLY the required columns, discard everything else (including Media Duration, etc.)
        normalized_data = []
        for row in data:
            normalized_row = {}
            for standard_name, actual_name in column_mapping.items():
                value = row.get(actual_name, '')
                # Ensure None values become empty strings
                normalized_row[standard_name] = value if value is not None else ''
            # Note: We deliberately do NOT keep other columns to avoid confusion
            normalized_data.append(normalized_row)

        return normalized_data

    def load_csv_data(self) -> bool:
        """Load and validate CSV data"""
        try:
            # Try different encodings to handle BOM and other issues
            csv_path = Path(self.csv_file)
            for encoding in SUPPORTED_ENCODINGS:
                try:
                    with csv_path.open('r', encoding=encoding) as file:
                        reader = csv.DictReader(file)
                        raw_data = list(reader)
                        break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            else:
                raise FileProcessingError("Could not decode CSV file with any supported encoding")

            # Clean and validate data
            cleaned_data = self._clean_column_names(raw_data)
            self.data = self._validate_and_normalize_columns(cleaned_data)
            return True

        except (CSVValidationError, FileProcessingError) as e:
            messagebox.showerror("Fel", f"Misslyckades ladda CSV-fil: {str(e)}")
            return False
        except Exception as e:
            messagebox.showerror("Fel", f"Oväntat fel vid laddning av CSV-fil: {str(e)}")
            return False

    def convert_to_word(self, output_path: str, filename: str, include_filename_prefix: bool = False) -> bool:
        """Convert CSV data to Word document"""
        try:
            if not DOCX_AVAILABLE:
                raise ConversionError("python-docx inte installerat. Kan inte skapa Word-dokument.\nInstallera med: pip install python-docx")

            doc = Document()

            # Add filename as header
            header_para = doc.add_heading(filename, 0)
            header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph()  # Empty line after header

            # Process each transcript entry
            for i, row in enumerate(self.data):
                timestamp = row['Media Start']
                transcript = (row.get('Transcript') or '').strip()
                speaker = (row.get('Speakers') or '').strip()

                # Create timecode with optional filename prefix
                if include_filename_prefix:
                    timecode_text = f"({filename}) [{timestamp}]"
                else:
                    timecode_text = f"[{timestamp}]"

                # Add timecode in square brackets (always, even for empty transcript)
                timecode_para = doc.add_paragraph()
                timecode_run = timecode_para.add_run(timecode_text)
                timecode_run.bold = True

                # Add transcript text with speaker if it exists
                if transcript:
                    if speaker:
                        # Include speaker name followed by colon
                        text_content = f"{speaker}: {transcript}"
                    else:
                        text_content = transcript
                    doc.add_paragraph(text_content)

                # Add extra paragraph for double line break before next timecode
                # (except for the last entry)
                if i < len(self.data) - 1:
                    doc.add_paragraph()

            # Save document
            output_file = Path(output_path)
            doc.save(str(output_file))
            return True

        except ConversionError as e:
            messagebox.showerror("Fel", str(e))
            return False
        except Exception as e:
            messagebox.showerror("Fel", f"Oväntat fel vid skapande av Word-dokument: {str(e)}")
            return False

    def convert_to_inqscribe(self, output_path: str, filename: str, include_filename_prefix: bool = False) -> bool:
        """Convert CSV data to Inqscribe format"""
        try:
            # Build the text content with proper formatting
            text_content = ""

            for row in self.data:
                timestamp_seconds = self.parse_timecode(row['Media Start'])
                timestamp_formatted = self.seconds_to_timecode(timestamp_seconds).strip('[]')
                transcript = (row.get('Transcript') or '').strip()
                speaker = (row.get('Speakers') or '').strip()

                # Create timecode with optional filename prefix
                if include_filename_prefix:
                    timecode_text = f"({filename}) [{timestamp_formatted}]"
                else:
                    timecode_text = f"[{timestamp_formatted}]"

                # Always add timecode, even for empty transcript (to show time adjustments)
                text_content += f"\\r\\r{timecode_text}"
                if transcript:
                    if speaker:
                        # Include speaker name followed by colon
                        text_content += f": {speaker}: {transcript}"
                    else:
                        text_content += f": {transcript}"

            # Inqscribe file structure
            inqscribe_content = f"""app=InqScribe
font.name=Tahoma
font.size=12
print.bottom=1.
print.left=1.
print.right=1.
print.top=1.
print.units=1
state.aspectratio=0.
tc.format=[x]
tc.includesourcename=0
tc.omitframes=0
tc.unbracketed=0
text={filename}{text_content}
timecode.fps={FRAME_RATE}
type=none
version=1.1
warned.fpsconflict=0"""

            # Write to file
            output_file = Path(output_path)
            output_file.write_text(inqscribe_content, encoding='utf-8')
            return True

        except FileProcessingError as e:
            messagebox.showerror("Fel", str(e))
            return False
        except Exception as e:
            messagebox.showerror("Fel", f"Oväntat fel vid skapande av Inqscribe-fil: {str(e)}")
            return False


class LimecraftGUI:
    def __init__(self):
        self.converter = LimecraftConverter()
        self.setup_gui()

    def _set_window_icon(self, window):
        """Set the egg icon for a window"""
        try:
            # Try to find the egg icon in the same folder as the script
            script_dir = Path(__file__).parent
            icon_path = script_dir / "egg_icon.png"

            if icon_path.exists():
                # Load the icon image
                from PIL import Image, ImageTk
                icon_image = Image.open(icon_path)
                # Resize to standard icon size
                icon_image = icon_image.resize((32, 32), Image.Resampling.LANCZOS)
                icon_photo = ImageTk.PhotoImage(icon_image)

                # Set the icon
                window.iconphoto(True, icon_photo)
                # Keep a reference to prevent garbage collection
                window.icon_photo = icon_photo
        except Exception:
            # If icon loading fails, continue without icon
            pass

    def setup_gui(self) -> None:
        """Setup the GUI interface"""
        # Create main window with superhero theme
        self.root = tb.Window(themename="superhero")
        self.root.title("Limecraft CSV-konverterare")

        # Set the window icon
        self._set_window_icon(self.root)

        # Get screen dimensions and set appropriate window size
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # Calculate window size as percentage of screen (max 900x750, min 700x580)
        window_width = min(900, max(700, int(screen_width * 0.7)))
        window_height = min(750, max(580, int(screen_height * 0.8)))

        # Center the window
        x = (screen_width // 2) - (window_width // 2)
        y = (screen_height // 2) - (window_height // 2)

        self.root.geometry(f"{window_width}x{window_height}+{x}+{y}")
        self.root.minsize(700, 580)  # Set minimum window size (increased height)

        # Variables
        self.csv_file_var = tk.StringVar(value="Ingen CSV-fil vald")
        self.start_time_var = tk.StringVar(value="")
        self.filename_var = tk.StringVar(value="")
        self.word_var = tk.BooleanVar()
        self.inqscribe_var = tk.BooleanVar()
        self.include_filename_var = tk.BooleanVar()  # New option for filename prefix

        self.create_widgets()

    def _create_help_content(self, main_frame) -> None:
        """Create help text content"""
        help_text = """Så här använder du programmet:

1. Välj CSV-fil (som du exporterats från Limecraft)
   Klicka på "Bläddra CSV-fil..." och välj den CSV-fil som du har exporterat från Limecraft.
   Programmet kommer att läsa in filen och visa dess namn i gränssnittet här nedanför.
   (För att exportera din transkribering som csv-fil från Limecraft välj "Export" och klicka på "CSV-fil".
   Rutorna framför "Media Start", "Transcript" och "Speakers" måste vara ikryssade. Det är ok om fler rutor är ikryssade men ingen annan metadata används.)

2. Justera starttid (valfritt)
   Ange tid som ska LÄGGAS TILL alla tidskoder. T ex starttiden på kamerakortet.
   Tidkoden du lägger till kommer att omvandlas till formatet HH:MM:SS.FF (timmar:minuter:sekunder.frames).
   Men det går bra att mata in tidkoden i de här formaten:
   • HH:MM:SS:FF (t.ex. 01:30:45:12)
   • HH:MM:SS.FF (t.ex. 01:30:45.12)
   • HH.MM.SS.FF (t.ex. 01.30.45.12)
   • HHMMSSFF (t.ex. 01304512)

3. Välj filnamn (valfritt)
   Ange önskat filnamn (utan filnamnstillägg).
   Samma namn kommer att användas för både .docx- och .inqscr-filer.
   Om du inte skriver in något namn används filnamnet från den valda CSV-filen (men med nya filnamnstillägg).

4. Välj utdataformat
   Markera om du vill ha ut transkriberingen som Word-dokument (.docx) och/eller som en Inqscribe-fil (.inqscr).

5. Välj om du vill lägga till filnamnet före varje tidskod
Du kan här välja att lägga till filnamnet (som står i rutan ovanför) inom parentes före varje tidskod:
(Filnamn) [HH:MM:SS.FF]
Exempel: Om ditt filnamn är "2025-06-07 synk Teet Härm" så kommer alla tidkoder i transkriberingsfilen se ut så här:
(2025-06-07 synk Teet Härm) [HH:MM:SS.FF]
Detta kan vara en fördel om du lägger in bitar från många olika intervjuer i ditt manus och vill att t ex klipparen lätt ska kunna se varifrån alla klipp kommer.

5. Konvertera
   Klicka på "Konvertera filer". Du får då välja en mapp där filerna ska sparas.

6. Klart!

Programmet lägger automatiskt till [00:00:00.00] i början av transkriptionen. Alternativt den nya starttid som du själv valt.
Programmet inkluderar nu även talarnamn från Speakers-kolumnen i utdatafilerna.

Klicka på Ett ägg med smör i-knappen för att stänga hjälpfönstret.
"""

        # Scrollable text widget
        text_frame = tb.Frame(main_frame)
        text_frame.pack(fill="both", expand=True, pady=(0, 20))

        text_widget = scrolledtext.ScrolledText(text_frame, wrap=tk.WORD,
                                               font=("Arial", 10), height=12)
        text_widget.pack(fill="both", expand=True)
        text_widget.insert(tk.END, help_text)
        text_widget.config(state=tk.DISABLED)

    def _create_help_image(self, bottom_frame) -> None:
        """Create help window image (egg with butter)"""
        try:
            from PIL import Image, ImageTk

            # Try to find the egg image in the same folder as the script
            script_dir = Path(__file__).parent
            egg_image_path = script_dir / "Agg_med_smor_v4_transperant.png"

            if egg_image_path.exists():
                # Load and display the actual PNG image
                pil_image = Image.open(egg_image_path)

                # Resize to appropriate size for help window
                pil_image = pil_image.resize((120, 120), Image.Resampling.LANCZOS)

                # Convert to PhotoImage
                egg_photo = ImageTk.PhotoImage(pil_image)

                # Keep a reference to prevent garbage collection
                bottom_frame.egg_image = egg_photo

                # Display the image
                image_label = tb.Label(bottom_frame, image=egg_photo)
                image_label.pack(pady=(10, 20))

            else:
                # PNG not found, use emoji fallback
                image_label = tb.Label(bottom_frame, text="🥚", font=("Arial", 72))
                image_label.pack(pady=(10, 10))

                butter_label = tb.Label(bottom_frame, text="🧈", font=("Arial", 48))
                butter_label.pack(pady=(0, 20))

        except Exception:
            # Enhanced fallback with better formatting
            fallback_frame = tb.Frame(bottom_frame)
            fallback_frame.pack(pady=(10, 20))

            tb.Label(fallback_frame, text="🥚", font=("Arial", 72)).pack()
            tb.Label(fallback_frame, text="🧈", font=("Arial", 48)).pack(pady=(10, 0))
            tb.Label(fallback_frame, text="Ett ägg med smör",
                     font=("Arial", 12, "italic")).pack(pady=(10, 0))

    def show_help(self) -> None:
        """Show help window with egg image"""
        help_win = tb.Toplevel(self.root)
        help_win.title("Hjälp")
        help_win.geometry("600x900")

        # Set the icon for help window too
        self._set_window_icon(help_win)

        # Center the help window at the same height as the main window
        help_win.update_idletasks()
        main_x = self.root.winfo_x()
        main_y = self.root.winfo_y()
        main_width = self.root.winfo_width()

        # Position help window to the right of main window, same Y position
        x = main_x + main_width + 20  # 20px gap between windows
        y = main_y  # Same height as main window

        help_win.geometry(f"600x900+{x}+{y}")

        # Make window appear on top
        help_win.lift()
        help_win.focus_force()
        help_win.attributes('-topmost', True)
        help_win.after(100, lambda: help_win.attributes('-topmost', False))

        # Main frame
        main_frame = tb.Frame(help_win)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Header
        tb.Label(main_frame, text="Limecraft CSV-konverterare - Hjälp",
                 font=("Arial", 14, "bold")).pack(pady=(0, 15))

        # Create help content
        self._create_help_content(main_frame)

        # Create a frame for the image and button with more space
        bottom_frame = tb.Frame(main_frame)
        bottom_frame.pack(pady=(20, 10))

        # Create image
        self._create_help_image(bottom_frame)

        # Close button with the internal joke
        close_btn = tb.Button(bottom_frame, text="Ett ägg med smör i",
                             command=help_win.destroy, bootstyle=DANGER, width=25)
        close_btn.pack()

    def create_widgets(self) -> None:
        """Create and layout GUI widgets"""
        # Create outer container for canvas and scrollbar
        outer_frame = tb.Frame(self.root)
        outer_frame.pack(fill="both", expand=True)

        # Create a canvas and scrollbar for scrollable content
        canvas = tk.Canvas(outer_frame)
        scrollbar = tb.Scrollbar(outer_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = tb.Frame(canvas)

        # Configure scrolling
        def configure_scroll_region(event=None):
            canvas.configure(scrollregion=canvas.bbox("all"))

        def configure_canvas_width(event=None):
            # Make the canvas window fill the canvas width
            canvas_width = canvas.winfo_width()
            canvas.itemconfig(canvas_window, width=canvas_width)

        scrollable_frame.bind("<Configure>", configure_scroll_region)
        canvas.bind("<Configure>", configure_canvas_width)

        # Create the window inside canvas and store reference
        canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        # Pack canvas and scrollbar
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # Bind mousewheel to canvas
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")

        def bind_mousewheel(event):
            canvas.bind_all("<MouseWheel>", _on_mousewheel)

        def unbind_mousewheel(event):
            canvas.unbind_all("<MouseWheel>")

        canvas.bind('<Enter>', bind_mousewheel)
        canvas.bind('<Leave>', unbind_mousewheel)

        # Main container with reduced padding - now inside scrollable frame
        main_frame = tb.Frame(scrollable_frame)
        main_frame.pack(fill="both", expand=True, padx=20, pady=15)

        # Top frame with title and help button
        top_frame = tb.Frame(main_frame)
        top_frame.pack(fill="x", pady=(0, 10))

        # Title - smaller font
        title_label = tb.Label(top_frame, text="Limecraft CSV-konverterare",
                              font=('Arial', 14, 'bold'))
        title_label.pack(side="left")

        # Help button
        help_btn = tb.Button(top_frame, text="Hjälp", command=self.show_help,
                            bootstyle=INFO, width=8)
        help_btn.pack(side="right")

        # Subtitle - smaller font and reduced padding
        subtitle_label = tb.Label(main_frame,
                                 text="Konvertera transkriptionsfiler från Limecraft till Word- och Inqscribe-format",
                                 font=('Arial', 9), wraplength=700)
        subtitle_label.pack(pady=(0, 15))

        # File selection section - reduced padding
        file_section = tb.LabelFrame(main_frame, text="1. Välj CSV-fil", padding=12)
        file_section.pack(fill="x", pady=(0, 10))

        file_label = tb.Label(file_section, textvariable=self.csv_file_var,
                             foreground="lightgray", wraplength=700, font=('Arial', 9))
        file_label.pack(anchor="w", pady=(0, 8))

        select_btn = tb.Button(file_section, text="Bläddra CSV-fil...",
                              command=self.select_csv_file, bootstyle=PRIMARY, width=18)
        select_btn.pack(anchor="w")

        # Time adjustment section - reduced padding
        time_section = tb.LabelFrame(main_frame, text="2. Justera starttid (valfritt)", padding=12)
        time_section.pack(fill="x", pady=(0, 10))

        time_info = tb.Label(time_section,
                            text="Addera tid till alla tidskoder (format HH:MM:SS.FF):",
                            font=('Arial', 9), wraplength=700)
        time_info.pack(anchor="w", pady=(0, 8))

        time_frame = tb.Frame(time_section)
        time_frame.pack(fill="x")

        tb.Label(time_frame, text="Lägg till:", font=('Arial', 9)).pack(side="left")
        time_entry = tb.Entry(time_frame, textvariable=self.start_time_var,
                             width=15, font=('Consolas', 9))
        time_entry.pack(side="left", padx=(8, 0))

        example_label = tb.Label(time_frame, text="(01:08:18:13 eller 01081813)",
                                font=('Arial', 8), foreground="gray")
        example_label.pack(side="left", padx=(8, 0))

        # Output filename section - reduced padding
        filename_section = tb.LabelFrame(main_frame, text="3. Filnamn (valfritt)", padding=12)
        filename_section.pack(fill="x", pady=(0, 10))

        filename_info = tb.Label(filename_section,
                                text="Anpassat filnamn (utan tillägg) eller lämna blankt:",
                                font=('Arial', 9), wraplength=700)
        filename_info.pack(anchor="w", pady=(0, 8))

        filename_frame = tb.Frame(filename_section)
        filename_frame.pack(fill="x")

        tb.Label(filename_frame, text="Filnamn:", font=('Arial', 9)).pack(side="left")
        filename_entry = tb.Entry(filename_frame, textvariable=self.filename_var,
                                 width=40, font=('Arial', 9))
        filename_entry.pack(side="left", padx=(8, 0), fill="x", expand=True)

        # Output format section - reduced padding
        format_section = tb.LabelFrame(main_frame, text="4. Utdataformat", padding=12)
        format_section.pack(fill="x", pady=(0, 10))

        word_check = tb.Checkbutton(format_section, text="Word-dokument (.docx)",
                                   variable=self.word_var, bootstyle="info-round-toggle")
        word_check.pack(anchor="w", pady=3)

        if not DOCX_AVAILABLE:
            word_check.config(state="disabled")
            tb.Label(format_section, text="  ⚠ python-docx inte installerat",
                    foreground="orange", font=('Arial', 8)).pack(anchor="w")

        inqscribe_check = tb.Checkbutton(format_section, text="Inqscribe-fil (.inqscr)",
                                        variable=self.inqscribe_var, bootstyle="info-round-toggle")
        inqscribe_check.pack(anchor="w", pady=3)

        # Filename prefix option - compact text
        filename_prefix_check = tb.Checkbutton(format_section, text="Inkludera filnamn före tidskoder",
                                              variable=self.include_filename_var, bootstyle="info-round-toggle")
        filename_prefix_check.pack(anchor="w", pady=3)

        # Action buttons section - reduced padding
        action_section = tb.Frame(main_frame)
        action_section.pack(fill="x", pady=(10, 0))

        convert_btn = tb.Button(action_section, text="Konvertera filer",
                               command=self.convert_files, bootstyle=SUCCESS, width=25)
        convert_btn.pack(pady=(0, 8))

        # Progress bar (initially hidden) - reduced padding
        self.progress_frame = tb.Frame(action_section)
        self.progress_label = tb.Label(self.progress_frame, text="", font=("Arial", 9))
        self.progress_label.pack(pady=(3, 0))

        self.progress_bar = tb.Progressbar(self.progress_frame, mode='indeterminate',
                                          bootstyle="success-striped")
        self.progress_bar.pack(fill="x", pady=(0, 3))

        # Status label - reduced padding
        self.status_var = tk.StringVar(value="Redo")
        status_label = tb.Label(action_section, textvariable=self.status_var,
                               font=('Arial', 9))
        status_label.pack(pady=(5, 10))

        # Version label in lower right corner of main window
        version_label = tb.Label(self.root, text=f"v{VERSION}",
                                font=('Arial', 7), foreground="gray")
        version_label.place(relx=1.0, rely=1.0, anchor="se", x=-10, y=-10)

    def select_csv_file(self) -> None:
        """Handle CSV file selection"""
        file_path = filedialog.askopenfilename(
            title="Välj Limecraft CSV-fil",
            filetypes=[("CSV-filer", "*.csv"), ("Alla filer", "*.*")]
        )

        if file_path:
            self.converter.csv_file = file_path
            filename = Path(file_path).name
            # Auto-populate filename field with CSV name (without extension)
            base_filename = Path(file_path).stem
            self.filename_var.set(base_filename)
            self.csv_file_var.set(f"Vald: {filename}")
            self.status_var.set("CSV-fil vald")

    def validate_inputs(self) -> bool:
        """Validate user inputs"""
        if not self.converter.csv_file:
            messagebox.showwarning("Varning", "Välj en CSV-fil först.")
            return False

        if not self.filename_var.get().strip():
            messagebox.showwarning("Varning", "Ange ett filnamn.")
            return False

        if not self.word_var.get() and not self.inqscribe_var.get():
            messagebox.showwarning("Varning", "Välj minst ett utdataformat.")
            return False

        # Validate time format if provided
        if self.start_time_var.get().strip():
            try:
                # Test normalize the time input
                self.converter.normalize_time_input(self.start_time_var.get().strip())
            except TimecodeFormatError as e:
                messagebox.showerror("Fel", f"Ogiltigt tidsformat: {str(e)}\n\nAccepterade format:\n• HH:MM:SS.FF\n• HH.MM.SS.FF\n• HHMMSSFF")
                return False

        return True

    def _perform_conversions(self, output_dir: str, custom_filename: str, filename_header: str) -> tuple[int, list[str]]:
        """Perform the actual file conversions"""
        success_count = 0
        created_files = []
        output_dir_path = Path(output_dir)
        include_filename = self.include_filename_var.get()

        # Convert to Word if selected
        if self.word_var.get():
            self.progress_label.config(text="Konverterar till Word...")
            self.root.update()
            word_path = output_dir_path / f"{custom_filename}.docx"
            if self.converter.convert_to_word(str(word_path), filename_header, include_filename):
                success_count += 1
                created_files.append(f"Word: {word_path.name}")

        # Convert to Inqscribe if selected
        if self.inqscribe_var.get():
            self.progress_label.config(text="Konverterar till Inqscribe...")
            self.root.update()
            inqscribe_path = output_dir_path / f"{custom_filename}.inqscr"
            if self.converter.convert_to_inqscribe(str(inqscribe_path), filename_header, include_filename):
                success_count += 1
                created_files.append(f"Inqscribe: {inqscribe_path.name}")

        return success_count, created_files

    def _show_conversion_results(self, success_count: int, created_files: list[str]) -> None:
        """Show conversion results to user"""
        if success_count > 0:
            files_text = "\n".join(f"• {file}" for file in created_files)
            messagebox.showinfo("Lyckades",
                               f"Konvertering slutförd framgångsrikt!\n\n"
                               f"Skapade {success_count} fil(er):\n{files_text}")
            self.status_var.set(f"Konvertering slutförd - {success_count} fil(er) skapade")
        else:
            self.status_var.set("Konvertering misslyckades")

    def convert_files(self) -> None:
        """Handle the conversion process"""
        if not self.validate_inputs():
            return

        # Show progress
        self.progress_frame.pack(pady=(10, 0), fill="x")
        self.progress_label.config(text="Laddar CSV-data...")
        self.progress_bar.start()
        self.root.update()

        try:
            # Load CSV data
            if not self.converter.load_csv_data():
                self.progress_bar.stop()
                self.progress_frame.pack_forget()
                self.status_var.set("Misslyckades ladda CSV")
                return

            # Add starting timecode [00:00:00:00] if needed
            # Check if user will apply time adjustment
            will_adjust_time = bool(self.start_time_var.get().strip())
            self.converter.add_starting_timecode(time_adjustment_applied=will_adjust_time)

            # Adjust timecodes if requested (add time to all timecodes)
            if self.start_time_var.get().strip():
                self.progress_label.config(text="Justerar tidskoder...")
                self.root.update()
                try:
                    self.converter.adjust_timecodes(self.start_time_var.get().strip())
                except TimecodeFormatError as e:
                    self.progress_bar.stop()
                    self.progress_frame.pack_forget()
                    messagebox.showerror("Tidsjusteringsfel", str(e))
                    return

            # Get output directory
            self.progress_bar.stop()
            self.progress_frame.pack_forget()

            output_dir = filedialog.askdirectory(title="Välj utdatamapp")
            if not output_dir:
                self.status_var.set("Avbruten")
                return

            # Show progress again for conversion
            self.progress_frame.pack(pady=(10, 0), fill="x")
            self.progress_bar.start()

            # Generate output filenames using custom filename
            custom_filename = self.filename_var.get().strip()
            # Remove any file extensions if user accidentally included them
            custom_filename = custom_filename.replace('.docx', '').replace('.inqscr', '')
            filename_header = custom_filename  # Use custom filename as header

            # Perform conversions
            success_count, created_files = self._perform_conversions(output_dir, custom_filename, filename_header)

            self.progress_bar.stop()
            self.progress_frame.pack_forget()

            # Show results
            self._show_conversion_results(success_count, created_files)

        except Exception as e:
            self.progress_bar.stop()
            self.progress_frame.pack_forget()
            messagebox.showerror("Fel", f"Ett fel uppstod under konverteringen: {str(e)}")
            self.status_var.set("Konvertering misslyckades")

    def run(self) -> None:
        """Start the GUI"""
        self.root.mainloop()


def main() -> None:
    """Main function"""
    # Check for required packages
    missing_packages = []

    try:
        import ttkbootstrap  # noqa: F401
    except ImportError:
        missing_packages.append("ttkbootstrap")

    if missing_packages:
        print("Saknade nödvändiga paket:")
        for package in missing_packages:
            print(f"  pip install {package}")
        return

    if not DOCX_AVAILABLE:
        print("Varning: python-docx inte installerat. Word-export kommer inte vara tillgängligt.")
        print("Installera med: pip install python-docx")

    # Start GUI
    app = LimecraftGUI()
    app.run()


if __name__ == "__main__":
    main()